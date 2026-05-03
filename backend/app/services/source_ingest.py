"""上传源材料解析服务。"""
from __future__ import annotations

import os
import posixpath
import re
import shutil
import zipfile
from pathlib import Path
from typing import Any


QUESTION_SOURCE_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
PAPER_SOURCE_EXTENSIONS = {".zip", ".md", ".docx", ".pdf"}
SUPPLEMENTARY_EXTENSIONS = {
    ".csv",
    ".xlsx",
    ".xls",
    ".json",
    ".txt",
    ".dat",
    ".tsv",
    ".md",
    ".png",
    ".jpg",
    ".jpeg",
    ".svg",
    ".webp",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".svg", ".webp", ".gif", ".bmp"}
MARKDOWN_IMAGE_PATTERN = re.compile(r"!\[[^\]]*\]\(([^)]+)\)")


class SourceIngestError(RuntimeError):
    """源材料解析失败。"""


def _safe_extract_zip(zip_path: str, dest_dir: str) -> None:
    dest_root = Path(dest_dir).resolve()
    with zipfile.ZipFile(zip_path, "r") as archive:
        for member in archive.infolist():
            member_path = (dest_root / member.filename).resolve()
            if not str(member_path).startswith(str(dest_root)):
                raise SourceIngestError("压缩包包含非法路径")
        archive.extractall(dest_root)


def _extract_pdf_text(filepath: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise SourceIngestError("缺少 pypdf，无法解析 PDF") from exc

    try:
        reader = PdfReader(filepath)
        parts = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as exc:
        raise SourceIngestError(f"PDF 解析失败: {exc}") from exc

    text = "\n\n".join(part for part in parts if part)
    if not text.strip():
        raise SourceIngestError("PDF 未提取到可用文本")
    return text


def _extract_docx_text(filepath: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise SourceIngestError("缺少 python-docx，无法解析 DOCX") from exc

    try:
        document = Document(filepath)
    except Exception as exc:
        raise SourceIngestError(f"DOCX 解析失败: {exc}") from exc

    lines: list[str] = []
    for paragraph in document.paragraphs:
        content = paragraph.text.strip()
        if content:
            lines.append(content)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    text = "\n".join(lines).strip()
    if not text:
        raise SourceIngestError("DOCX 未提取到可用文本")
    return text


def _read_text_like_file(filepath: str) -> str:
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf_text(filepath)
    if ext == ".docx":
        return _extract_docx_text(filepath)

    try:
        return Path(filepath).read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return Path(filepath).read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        raise SourceIngestError(f"文件读取失败: {exc}") from exc


def _pick_primary_markdown(bundle_dir: Path) -> Path:
    preferred_names = ["res.md", "paper.md", "main.md", "index.md", "README.md"]
    markdown_files = sorted(bundle_dir.rglob("*.md"))
    if not markdown_files:
        raise SourceIngestError("压缩包中未找到 Markdown 文档")

    for preferred in preferred_names:
        for candidate in markdown_files:
            if candidate.name.lower() == preferred.lower():
                return candidate
    return markdown_files[0]


def _normalize_markdown_refs(markdown_text: str, markdown_path: Path, bundle_dir: Path, bundle_prefix: str) -> tuple[str, list[str]]:
    referenced_assets: list[str] = []

    def replace(match: re.Match[str]) -> str:
        raw_path = match.group(1).strip()
        if not raw_path or raw_path.startswith(("http://", "https://", "data:", "#")):
            return match.group(0)

        candidate = raw_path.split("?", 1)[0].split("#", 1)[0].strip()
        resolved = (markdown_path.parent / candidate).resolve()
        try:
            relative_in_bundle = resolved.relative_to(bundle_dir.resolve()).as_posix()
        except ValueError:
            return match.group(0)

        normalized = posixpath.join(bundle_prefix, relative_in_bundle)
        referenced_assets.append(normalized)
        return match.group(0).replace(raw_path, normalized)

    normalized_text = MARKDOWN_IMAGE_PATTERN.sub(replace, markdown_text)
    return normalized_text, referenced_assets


def _build_image_manifest(bundle_dir: Path, bundle_prefix: str, referenced_assets: set[str]) -> list[dict[str, Any]]:
    manifest: list[dict[str, Any]] = []
    for image_path in sorted(bundle_dir.rglob("*")):
        if not image_path.is_file() or image_path.suffix.lower() not in IMAGE_EXTENSIONS:
            continue

        relative_path = image_path.relative_to(bundle_dir).as_posix()
        normalized = posixpath.join(bundle_prefix, relative_path)
        manifest.append(
            {
                "path": normalized,
                "filename": image_path.name,
                "size_kb": round(image_path.stat().st_size / 1024, 1),
                "referenced": normalized in referenced_assets,
            }
        )
    return manifest


def ingest_question_source(saved_path: str, task: dict) -> dict[str, Any]:
    text = _read_text_like_file(saved_path).strip()
    updates = {
        "source_question_text": text,
        "source_question_file": saved_path,
    }
    if not (task.get("question") or "").strip():
        updates["question"] = text
    if not (task.get("source_question") or "").strip():
        updates["source_question"] = text
    return updates


def ingest_paper_source(saved_path: str, task: dict) -> dict[str, Any]:
    ext = Path(saved_path).suffix.lower()
    if ext == ".zip":
        bundle_dir = Path(task["work_dir"]) / "source_bundle"
        if bundle_dir.exists():
            shutil.rmtree(bundle_dir)
        bundle_dir.mkdir(parents=True, exist_ok=True)

        _safe_extract_zip(saved_path, str(bundle_dir))
        markdown_path = _pick_primary_markdown(bundle_dir)
        raw_markdown = _read_text_like_file(str(markdown_path)).strip()
        normalized_markdown, referenced_assets = _normalize_markdown_refs(
            raw_markdown,
            markdown_path,
            bundle_dir,
            "source_bundle",
        )
        manifest = _build_image_manifest(bundle_dir, "source_bundle", set(referenced_assets))
        return {
            "paper_content": normalized_markdown,
            "source_paper_file": saved_path,
            "source_bundle_dir": str(bundle_dir),
            "source_markdown_file": str(markdown_path),
            "source_image_manifest": manifest,
        }

    paper_text = _read_text_like_file(saved_path).strip()
    updates = {
        "paper_content": paper_text,
        "source_paper_file": saved_path,
        "source_bundle_dir": "",
        "source_markdown_file": saved_path if ext == ".md" else "",
        "source_image_manifest": [],
    }
    return updates


def purpose_for_upload(task_type: str, filename: str, declared_purpose: str = "auto") -> str:
    if declared_purpose and declared_purpose != "auto":
        return declared_purpose

    ext = Path(filename or "").suffix.lower()
    if task_type == "writing" and ext in QUESTION_SOURCE_EXTENSIONS:
        return "question_source"
    if task_type == "polish" and ext in PAPER_SOURCE_EXTENSIONS:
        return "paper_source"
    return "supplementary_data"