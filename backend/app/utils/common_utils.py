"""
通用工具函数
"""
import os
import json
import re
from pathlib import Path
from typing import Any, Optional


PLACEHOLDER_PATTERNS = [
    "张三", "李四", "王五", "赵六", "陈七", "刘八",
    "某某", "xxx", "xx", "todo", "tbd", "待补充", "参考文献待完善", "作者待补", "年份待补",
]

BAD_MATH_PATTERNS = [
    r"=\s*$",
    r"\(\s*\)",
    r"\[\s*\]",
    r"\$\$\s*\$\$",
    r"其中，\s*为",
]

UNIT_RULES = {
    "MPa": "N/mm^2",
    "kN*mm": "N*m",
    "GPa_to_MPa": 1000,
    "P_kN_to_N": 1000,
}

STRONG_RESULT_KEYS = {
    "id",
    "name",
    "value",
    "computed_value",
    "paper_value",
    "verified",
    "status",
    "formula",
    "inputs",
    "source_data",
    "code_cell",
}

BLOCKED_STATUS_HINTS = {"blocked", "failed", "mismatch", "unverified", "阻断", "失败", "不通过"}

LEGACY_RESULT_CONTEXT_KEYS = {
    "title",
    "name",
    "case",
    "problem_id",
    "工况",
    "场景",
    "项目",
    "对象",
}

LEGACY_RESULT_IGNORED_KEYS = {
    "title",
    "name",
    "case",
    "problem_id",
    "工况",
    "场景",
    "项目",
    "对象",
    "说明",
    "描述",
    "建议",
    "结论",
    "限制因素",
}


def _slugify_identifier(text: str, fallback: str = "item") -> str:
    value = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff_]+", "_", (text or "").strip())
    value = re.sub(r"_+", "_", value).strip("_")
    return value or fallback


def _parse_numeric_value(raw_value: Any) -> float | None:
    if isinstance(raw_value, (int, float)):
        return float(raw_value)
    if raw_value is None:
        return None
    text = str(raw_value).strip().replace(",", "")
    if not text:
        return None
    suffix = 0.01 if text.endswith("%") else 1.0
    if suffix != 1.0:
        text = text[:-1]
    try:
        return float(text) * suffix
    except Exception:
        return None


def _deduplicate_dict_rows(rows: list[dict[str, Any]], keys: list[str]) -> list[dict[str, Any]]:
    unique_rows: list[dict[str, Any]] = []
    seen: set[tuple[Any, ...]] = set()
    for row in rows:
        marker = tuple(row.get(key) for key in keys)
        if marker in seen:
            continue
        seen.add(marker)
        unique_rows.append(row)
    return unique_rows


def extract_symbol_assignments(text: str) -> dict[str, dict[str, Any]]:
    assignments: dict[str, dict[str, Any]] = {}
    pattern = re.compile(
        r"(?P<name>[A-Za-zΑ-Ωα-ω_][A-Za-z0-9_Α-Ωα-ω]*)\s*(?:=|：|:)\s*(?P<value>[-+]?\d+(?:\.\d+)?(?:[eE][-+]?\d+)?%?)(?:\s*(?P<unit>[A-Za-z%°/·*^\-\d\.]+))?"
    )

    for line_number, raw_line in enumerate((text or "").splitlines(), start=1):
        for match in pattern.finditer(raw_line):
            name = match.group("name")
            assignments[name] = {
                "name": name,
                "value": match.group("value"),
                "numeric_value": _parse_numeric_value(match.group("value")),
                "unit": (match.group("unit") or "").strip(),
                "line": line_number,
                "source": raw_line.strip(),
            }
    return assignments


def extract_locked_parameters(question_text: str) -> dict[str, dict[str, Any]]:
    locked = extract_symbol_assignments(question_text)
    for item in locked.values():
        item["can_override"] = False
    return locked


def infer_problem_categories(question_text: str, breakdown_text: str = "") -> list[str]:
    text = f"{question_text}\n{breakdown_text}".lower()
    category_rules = [
        ("数据拟合型", ["拟合", "回归", "curve fit", "regression"]),
        ("优化决策型", ["最优", "优化", "决策", "规划"]),
        ("物理机理型", ["力", "力矩", "应力", "应变", "物理", "机理", "锚杆", "围岩"]),
        ("评价决策型", ["评价", "评分", "指标体系", "层次分析"]),
        ("预测型", ["预测", "forecast", "时间序列"]),
        ("图论/路径型", ["路径", "最短路", "网络", "图论"]),
        ("微分方程型", ["微分方程", "偏微分", "ode", "pde"]),
        ("多目标规划型", ["多目标", "pareto", "权衡"]),
        ("参数敏感性分析型", ["敏感性", "灵敏度", "扰动分析"]),
    ]
    categories = [name for name, keywords in category_rules if any(keyword in text for keyword in keywords)]
    return categories or ["通用建模型"]


def extract_breakdown_tasks(breakdown_text: str) -> list[dict[str, Any]]:
    tasks: list[dict[str, Any]] = []
    for line in (breakdown_text or "").splitlines():
        stripped = line.strip()
        if re.match(r"^(\d+\.|[-*])\s+", stripped):
            tasks.append({"text": re.sub(r"^(\d+\.|[-*])\s+", "", stripped), "source": stripped})
    return tasks[:12]


def build_problem_facts(question_text: str, breakdown_text: str) -> dict[str, Any]:
    locked_parameters = extract_locked_parameters(question_text)
    parameter_lines = extract_problem_parameters(question_text)
    data_tables = []
    for line in (question_text or "").splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if any(keyword in stripped for keyword in ["表", "数据", "附件", "附录"]):
            data_tables.append({"text": stripped})

    given_formulas = extract_markdown_formulas(question_text)
    tasks = extract_breakdown_tasks(breakdown_text)
    problem_types = infer_problem_categories(question_text, breakdown_text)

    return {
        "problem_types": problem_types,
        "constants": [
            {
                "name": name,
                "value": item.get("value", ""),
                "unit": item.get("unit", ""),
                "source": item.get("source", ""),
                "can_override": False,
            }
            for name, item in locked_parameters.items()
        ],
        "given_formulas": given_formulas,
        "data_tables": data_tables,
        "tasks": tasks,
        "forbidden_changes": [
            {
                "name": name,
                "locked_value": item.get("value", ""),
                "unit": item.get("unit", ""),
                "reason": "题设显式给定，后续阶段不得擅自修改",
            }
            for name, item in locked_parameters.items()
        ],
        "locked_parameters": locked_parameters,
        "parameter_lines": parameter_lines,
    }


def _normalize_result_warnings(raw_warnings: Any) -> list[str]:
    if isinstance(raw_warnings, list):
        return [str(item).strip() for item in raw_warnings if str(item).strip()]
    if raw_warnings is None:
        return []
    text = str(raw_warnings).strip()
    return [text] if text else []


def _looks_like_strong_result_entry(entry: Any) -> bool:
    return isinstance(entry, dict) and any(key in entry for key in STRONG_RESULT_KEYS)


def _normalize_legacy_scalar_result(
    *,
    section: str,
    source_file: str,
    path_parts: list[str],
    raw_value: Any,
    unit: str = "",
    status: str = "unverified",
    verified: bool = False,
    warnings: list[str] | None = None,
    source: str = "legacy_nested_results",
) -> dict[str, Any]:
    label = " / ".join(part for part in path_parts if part)
    text_value = str(raw_value).strip() if raw_value is not None else ""
    numeric_value = _parse_numeric_value(raw_value)
    return {
        "id": _slugify_identifier(f"{section}_{label}", "result"),
        "name": label or f"{section}_result",
        "section": section,
        "value": numeric_value if numeric_value is not None else text_value,
        "paper_value": text_value,
        "unit": unit,
        "formula": "",
        "formula_id": "",
        "inputs": {},
        "unit_conversion": "",
        "source_data": [],
        "code_cell": "",
        "evidence": f"{label}: {text_value}" if label else text_value,
        "source": source,
        "source_file": source_file,
        "generated_files": [],
        "status": status,
        "verified": verified,
        "warnings": warnings or [],
    }


def _split_legacy_metric_key(raw_key: str) -> tuple[str, str]:
    key = str(raw_key or "").strip()
    if not key:
        return "result", ""

    match = re.match(r"^(?P<label>.+?)_(?P<unit>[A-Za-z%°/·*^\-\d]+)$", key)
    if not match:
        return key, ""

    return match.group("label").strip() or key, match.group("unit").strip()


def _extract_legacy_key_result_metrics(
    item: dict[str, Any],
    *,
    section: str,
    source_file: str,
    is_verification_payload: bool,
) -> list[dict[str, Any]]:
    context = ""
    for key in LEGACY_RESULT_CONTEXT_KEYS:
        value = item.get(key)
        if value is not None and str(value).strip():
            context = str(value).strip()
            break

    entries: list[dict[str, Any]] = []
    for raw_key, raw_value in item.items():
        if raw_key in LEGACY_RESULT_IGNORED_KEYS:
            continue
        numeric_value = _parse_numeric_value(raw_value)
        if numeric_value is None:
            continue

        metric_name, metric_unit = _split_legacy_metric_key(str(raw_key))
        path_parts = [part for part in [context, metric_name] if part]
        entries.append(
            _normalize_legacy_scalar_result(
                section=section,
                source_file=source_file,
                path_parts=path_parts,
                raw_value=raw_value,
                unit=metric_unit,
                status="verified" if is_verification_payload else "verified",
                verified=True,
                warnings=["legacy_key_result_metric"],
                source="legacy_key_result_metric",
            )
        )

    return entries


def _collect_legacy_results_from_mapping(
    node: Any,
    *,
    section: str,
    source_file: str,
    path_parts: list[str],
    inherited_unit: str = "",
    status: str = "unverified",
    verified: bool = False,
    warnings: list[str] | None = None,
    source: str = "legacy_nested_results",
) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []

    if isinstance(node, dict):
        local_unit = inherited_unit
        if isinstance(node.get("unit"), (str, int, float)):
            local_unit = str(node.get("unit") or "").strip()

        for key, value in node.items():
            if key in {"unit", "description", "conclusion", "limitations", "applicability", "method"}:
                continue
            entries.extend(
                _collect_legacy_results_from_mapping(
                    value,
                    section=section,
                    source_file=source_file,
                    path_parts=path_parts + [str(key)],
                    inherited_unit=local_unit,
                    status=status,
                    verified=verified,
                    warnings=warnings,
                    source=source,
                )
            )
        return entries

    if isinstance(node, list):
        for index, item in enumerate(node, start=1):
            if isinstance(item, (dict, list)):
                entries.extend(
                    _collect_legacy_results_from_mapping(
                        item,
                        section=section,
                        source_file=source_file,
                        path_parts=path_parts + [str(index)],
                        inherited_unit=inherited_unit,
                        status=status,
                        verified=verified,
                        warnings=warnings,
                        source=source,
                    )
                )
                continue
            text = str(item).strip()
            if not text:
                continue
            entries.append(
                _normalize_legacy_scalar_result(
                    section=section,
                    source_file=source_file,
                    path_parts=path_parts + [str(index)],
                    raw_value=text,
                    unit=inherited_unit,
                    status=status,
                    verified=verified,
                    warnings=warnings,
                    source=source,
                )
            )
        return entries

    if node is None:
        return entries

    text = str(node).strip()
    if not text:
        return entries
    entries.append(
        _normalize_legacy_scalar_result(
            section=section,
            source_file=source_file,
            path_parts=path_parts,
            raw_value=node,
            unit=inherited_unit,
            status=status,
            verified=verified,
            warnings=warnings,
            source=source,
        )
    )
    return entries


def _extract_legacy_registry_entries(payload: dict[str, Any], section: str, source_file: str) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    verified_entries: list[dict[str, Any]] = []
    blocked_entries: list[dict[str, Any]] = []

    analysis_type = str(payload.get("analysis_type") or section)
    is_verification_payload = "复核" in analysis_type or analysis_type.lower() == "verification"

    problems = payload.get("problems")
    if isinstance(problems, dict):
        for problem_key, problem_payload in problems.items():
            if not isinstance(problem_payload, dict):
                continue
            problem_title = str(problem_payload.get("title") or problem_key)
            if isinstance(problem_payload.get("results"), dict):
                blocked_entries.extend(
                    _collect_legacy_results_from_mapping(
                        problem_payload.get("results"),
                        section=section,
                        source_file=source_file,
                        path_parts=[problem_title],
                        status="unverified",
                        verified=False,
                        warnings=["legacy_nested_result_requires_verification"],
                    )
                )
            subproblems = problem_payload.get("subproblems")
            if not isinstance(subproblems, dict):
                continue
            for sub_key, sub_payload in subproblems.items():
                if not isinstance(sub_payload, dict) or not isinstance(sub_payload.get("results"), dict):
                    continue
                sub_title = str(sub_payload.get("title") or sub_key)
                blocked_entries.extend(
                    _collect_legacy_results_from_mapping(
                        sub_payload.get("results"),
                        section=section,
                        source_file=source_file,
                        path_parts=[problem_title, sub_title],
                        status="unverified",
                        verified=False,
                        warnings=["legacy_nested_result_requires_verification"],
                    )
                )

    for item in payload.get("key_results", []):
        if _looks_like_strong_result_entry(item) or not isinstance(item, dict):
            continue
        metric_entries = _extract_legacy_key_result_metrics(
            item,
            section=section,
            source_file=source_file,
            is_verification_payload=is_verification_payload,
        )
        if metric_entries:
            verified_entries.extend(metric_entries)
        title = str(item.get("title") or item.get("problem_id") or f"{section}_review")
        findings = item.get("main_findings")
        if not isinstance(findings, list) or not findings:
            continue
        raw_status = str(item.get("verification_status") or payload.get("review_status") or "").strip().lower()
        is_blocked = raw_status in BLOCKED_STATUS_HINTS
        entry = _normalize_legacy_scalar_result(
            section=section,
            source_file=source_file,
            path_parts=[title],
            raw_value="；".join(str(finding).strip() for finding in findings if str(finding).strip()),
            status="blocked" if is_blocked else ("verified" if is_verification_payload else "unverified"),
            verified=(not is_blocked) and is_verification_payload,
            warnings=["legacy_review_summary"],
            source="legacy_review_key_results",
        )
        if is_blocked or not is_verification_payload:
            blocked_entries.append(entry)
        else:
            verified_entries.append(entry)

    for item in payload.get("blocking_items", []):
        if not isinstance(item, dict):
            continue
        label = str(item.get("item") or item.get("title") or "blocking_item")
        reason = str(item.get("reason") or item.get("recommendation") or "阻断项")
        blocked_entries.append(
            _normalize_legacy_scalar_result(
                section=section,
                source_file=source_file,
                path_parts=[label],
                raw_value=reason,
                status="blocked",
                verified=False,
                warnings=["blocking_item"],
                source="verification_blocking_item",
            )
        )

    for item in payload.get("prohibited_conclusions", []):
        text = str(item).strip()
        if not text:
            continue
        blocked_entries.append(
            _normalize_legacy_scalar_result(
                section=section,
                source_file=source_file,
                path_parts=["prohibited_conclusion"],
                raw_value=text,
                status="blocked",
                verified=False,
                warnings=["prohibited_conclusion"],
                source="verification_prohibited_conclusion",
            )
        )

    return verified_entries, blocked_entries


def _normalize_legacy_result_text(entry: str, section: str, source_file: str) -> dict[str, Any]:
    text = str(entry or "").strip()
    value_match = re.search(r"(?P<value>[-+]?\d+(?:\.\d+)?(?:[eE][-+]?\d+)?)\s*(?P<unit>[A-Za-z%°/·*^\-\d]+)?", text)
    formula_match = re.search(r"(?P<formula>[A-Za-zΑ-Ωα-ωPTEKσδfMRTLde_0-9\s\-+*/^().]+=[^，。；;]+)", text)
    blocked_hints = ["问题", "偏差", "不匹配", "无法", "失败", "警告", "风险", "待", "需", "需要", "谨慎"]
    verified = not any(hint in text for hint in blocked_hints)
    status = "verified" if verified else "blocked"
    name = text[:60] or f"{section}_result"

    return {
        "id": _slugify_identifier(f"{section}_{name}", "result"),
        "name": name,
        "section": section,
        "value": value_match.group("value") if value_match else "",
        "paper_value": "",
        "unit": (value_match.group("unit") or "").strip() if value_match else "",
        "formula": (formula_match.group("formula") or "").strip() if formula_match else "",
        "formula_id": "",
        "inputs": {},
        "unit_conversion": "",
        "source_data": [],
        "code_cell": "",
        "evidence": text,
        "source": "legacy_key_results",
        "source_file": source_file,
        "generated_files": [],
        "status": status,
        "verified": verified,
        "warnings": ["legacy_unstructured_result"] if verified else ["legacy_unstructured_result", "requires_structured_evidence"],
    }


def _normalize_registry_warning_entry(entry: str, section: str, source_file: str) -> dict[str, Any]:
    text = str(entry or "").strip()
    name = text[:60] or f"{section}_warning"
    return {
        "id": _slugify_identifier(f"{section}_warning_{name}", "warning"),
        "name": name,
        "section": section,
        "value": "",
        "paper_value": "",
        "unit": "",
        "formula": "",
        "formula_id": "",
        "inputs": {},
        "unit_conversion": "",
        "source_data": [],
        "code_cell": "",
        "evidence": text,
        "source": "structured_result_warning",
        "source_file": source_file,
        "generated_files": [],
        "status": "blocked",
        "verified": False,
        "warnings": ["reported_by_stage"],
    }


def normalize_result_registry_entry(entry: Any, section: str, source_file: str) -> dict[str, Any]:
    if isinstance(entry, str):
        return _normalize_legacy_result_text(entry, section, source_file)

    name = str(entry.get("name") or entry.get("id") or f"{section}_result").strip()
    entry_id = str(entry.get("id") or _slugify_identifier(f"{section}_{name}", "result")).strip()
    status = str(entry.get("status") or ("verified" if entry.get("verified", True) else "unverified")).strip().lower()
    warnings = _normalize_result_warnings(entry.get("warnings"))
    verified = bool(entry.get("verified", status not in {"mismatch", "blocked", "failed", "unverified"}))
    if warnings:
        verified = verified and all("无法确认" not in warning and "不可靠" not in warning for warning in warnings)

    return {
        "id": entry_id,
        "name": name,
        "section": section,
        "value": entry.get("value", entry.get("computed_value", "")),
        "paper_value": entry.get("paper_value", ""),
        "unit": entry.get("unit", ""),
        "formula": entry.get("formula", ""),
        "formula_id": entry.get("formula_id", ""),
        "inputs": entry.get("inputs", entry.get("input_values", {})) if isinstance(entry.get("inputs", entry.get("input_values", {})), dict) else {},
        "unit_conversion": entry.get("unit_conversion", ""),
        "source_data": entry.get("source_data", []),
        "code_cell": entry.get("code_cell", ""),
        "evidence": entry.get("evidence", ""),
        "source": entry.get("source", ""),
        "source_file": source_file,
        "generated_files": entry.get("generated_files", []),
        "status": status,
        "verified": verified,
        "warnings": warnings,
    }


def build_result_registry(work_dir: str, structured_result_files: list[str]) -> dict[str, Any]:
    verified_results: list[dict[str, Any]] = []
    blocked_results: list[dict[str, Any]] = []
    source_files: list[str] = []

    for filename in structured_result_files:
        if not filename:
            continue
        path = Path(work_dir) / filename
        payload = load_json(str(path))
        if not isinstance(payload, dict):
            continue
        source_files.append(filename)
        section = str(payload.get("section") or path.stem)
        for raw_entry in payload.get("key_results", []):
            if isinstance(raw_entry, dict) and not _looks_like_strong_result_entry(raw_entry):
                continue
            if not isinstance(raw_entry, (dict, str)):
                continue
            normalized = normalize_result_registry_entry(raw_entry, section, filename)
            if normalized["verified"] and normalized["status"] not in {"mismatch", "blocked", "failed"}:
                verified_results.append(normalized)
            else:
                blocked_results.append(normalized)
        extra_verified, extra_blocked = _extract_legacy_registry_entries(payload, section, filename)
        verified_results.extend(extra_verified)
        blocked_results.extend(extra_blocked)
        for warning in _normalize_result_warnings(payload.get("warnings")):
            blocked_results.append(_normalize_registry_warning_entry(warning, section, filename))

    verified_results = _deduplicate_dict_rows(verified_results, ["id", "source_file"])
    blocked_results = _deduplicate_dict_rows(blocked_results, ["id", "source_file"])

    return {
        "registry_version": "0.2",
        "source_files": source_files,
        "verified_results": verified_results,
        "blocked_results": blocked_results,
        "summary": {
            "verified_count": len(verified_results),
            "blocked_count": len(blocked_results),
        },
    }


def detect_placeholder_references(markdown_text: str, references: list[str]) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    lowered_patterns = [pattern.lower() for pattern in PLACEHOLDER_PATTERNS]
    for index, reference in enumerate(references, start=1):
        lower_reference = reference.lower()
        for pattern in lowered_patterns:
            if pattern in lower_reference:
                findings.append(
                    {
                        "type": "fake_reference",
                        "location": f"参考文献[{index}]",
                        "route": "writer",
                        "severity": "critical",
                        "fix": "删除或替换为真实检索得到的参考文献条目",
                        "evidence": reference,
                    }
                )
                break

    if not references and re.search(r"^#{1,6}\s*(参考文献|references)\s*$", markdown_text or "", re.I | re.M):
        findings.append(
            {
                "type": "fake_reference",
                "location": "参考文献章节",
                "route": "writer",
                "severity": "major",
                "fix": "若未使用外部文献，请删除参考文献章节或明确写未引入外部文献",
                "evidence": "存在参考文献标题但未提供真实条目",
            }
        )

    return findings


def detect_bad_math_issues(markdown_text: str) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    text = markdown_text or ""
    for pattern in BAD_MATH_PATTERNS:
        for match in re.finditer(pattern, text, re.M):
            line_number = text.count("\n", 0, match.start()) + 1
            findings.append(
                {
                    "type": "formula_missing",
                    "location": f"第{line_number}行附近",
                    "route": "writer",
                    "severity": "major",
                    "fix": "补全缺失或残缺的公式表达，并核对公式编号与变量定义",
                    "evidence": match.group(0),
                }
            )
    return _deduplicate_dict_rows(findings, ["type", "location", "evidence"])


def compare_locked_parameters(problem_facts: dict[str, Any], markdown_text: str) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    locked_parameters = (problem_facts or {}).get("locked_parameters", {})
    paper_assignments = extract_symbol_assignments(markdown_text)

    for name, locked in locked_parameters.items():
        paper_item = paper_assignments.get(name)
        if not paper_item:
            continue
        locked_value = locked.get("numeric_value")
        paper_value = paper_item.get("numeric_value")
        if locked_value is not None and paper_value is not None and abs(locked_value - paper_value) > 1e-9:
            findings.append(
                {
                    "type": "parameter_changed",
                    "location": f"参数 {name}",
                    "route": "modeling",
                    "severity": "critical",
                    "fix": f"恢复题设锁定参数 {name}={locked.get('value')} {locked.get('unit', '')}".strip(),
                    "evidence": f"题设为 {locked.get('value')}，论文中出现 {paper_item.get('value')}",
                }
            )
        locked_unit = str(locked.get("unit") or "").strip()
        paper_unit = str(paper_item.get("unit") or "").strip()
        if locked_unit and paper_unit and locked_unit != paper_unit:
            findings.append(
                {
                    "type": "unit_mismatch",
                    "location": f"参数 {name}",
                    "route": "modeling",
                    "severity": "major",
                    "fix": f"统一参数 {name} 的单位为 {locked_unit}",
                    "evidence": f"题设单位为 {locked_unit}，论文中出现 {paper_unit}",
                }
            )

    return _deduplicate_dict_rows(findings, ["type", "location", "evidence"])


def build_unit_audit_findings(problem_facts: dict[str, Any], result_registry: dict[str, Any]) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    locked_parameters = (problem_facts or {}).get("locked_parameters", {})
    registry_results = (result_registry or {}).get("verified_results", []) + (result_registry or {}).get("blocked_results", [])

    for entry in registry_results:
        if not str(entry.get("unit") or "").strip():
            findings.append(
                {
                    "type": "unit_missing",
                    "location": entry.get("name") or entry.get("id") or "未命名结果",
                    "route": "solver",
                    "severity": "major",
                    "fix": "为该结果补充明确单位，并在论文与表格中保持一致",
                    "evidence": f"结果 {entry.get('id', '')} 缺少单位字段",
                }
            )

        inputs = entry.get("inputs", {}) if isinstance(entry.get("inputs"), dict) else {}
        for param_name, input_item in inputs.items():
            if not isinstance(input_item, dict):
                continue
            locked = locked_parameters.get(param_name)
            if not locked:
                continue
            locked_unit = str(locked.get("unit") or "").strip()
            input_unit = str(input_item.get("unit") or "").strip()
            if locked_unit and input_unit and locked_unit != input_unit:
                findings.append(
                    {
                        "type": "unit_mismatch",
                        "location": entry.get("name") or entry.get("id") or param_name,
                        "route": "solver",
                        "severity": "major",
                        "fix": f"将输入参数 {param_name} 的单位统一为 {locked_unit}",
                        "evidence": f"题设单位 {locked_unit}，结果登记中使用 {input_unit}",
                    }
                )

    return _deduplicate_dict_rows(findings, ["type", "location", "evidence"])


def build_paper_audit_report(
    question_text: str,
    markdown_text: str,
    problem_facts: dict[str, Any],
    result_registry: dict[str, Any],
    audit_manifest: dict[str, Any],
    delivery_audit_text: str = "",
) -> dict[str, Any]:
    references = (audit_manifest or {}).get("references", [])
    blocks: list[dict[str, Any]] = []
    blocks.extend(detect_placeholder_references(markdown_text, references))
    blocks.extend(detect_bad_math_issues(markdown_text))
    blocks.extend(compare_locked_parameters(problem_facts, markdown_text))
    blocks.extend(build_unit_audit_findings(problem_facts, result_registry))

    for entry in (result_registry or {}).get("blocked_results", []):
        name = str(entry.get("name") or entry.get("id") or "").strip()
        if not name or name not in (markdown_text or ""):
            continue
        blocks.append(
            {
                "type": "numeric_mismatch",
                "location": name,
                "route": "solver",
                "severity": "critical",
                "fix": "重新计算该结果，并确保正文只引用 verified_results 中的数值",
                "evidence": f"结果登记状态为 {entry.get('status', 'unverified')}，但正文仍引用该结果",
            }
        )

    blocks = _deduplicate_dict_rows(blocks, ["type", "location", "evidence"])
    status = "BLOCK" if blocks else "PASS"
    return {
        "status": status,
        "blocks": blocks,
        "programmatic_checks": {
            "formula_count": (audit_manifest or {}).get("formula_count", 0),
            "table_count": (audit_manifest or {}).get("table_count", 0),
            "reference_count": (audit_manifest or {}).get("reference_count", 0),
            "locked_parameter_count": len((problem_facts or {}).get("locked_parameters", {})),
            "verified_result_count": len((result_registry or {}).get("verified_results", [])),
            "blocked_result_count": len((result_registry or {}).get("blocked_results", [])),
        },
        "llm_summary": (delivery_audit_text or "").strip(),
    }


def validate_paper_audit_report(payload: Any) -> bool:
    if not isinstance(payload, dict):
        return False
    if payload.get("status") not in {"PASS", "BLOCK"}:
        return False
    blocks = payload.get("blocks")
    if not isinstance(blocks, list):
        return False
    for block in blocks:
        if not isinstance(block, dict):
            return False
        required = ["type", "location", "route", "severity", "fix"]
        if any(not str(block.get(key, "")).strip() for key in required):
            return False
    return True


def _resolve_markdown_image_path(md_path: str, image_target: str) -> Path | None:
    target = (image_target or "").strip().strip("<>").strip()
    if not target or re.match(r"^[a-zA-Z]+://", target):
        return None

    if " \"" in target:
        target = target.split(' "', 1)[0].strip()

    md_dir = Path(md_path).parent
    raw_candidate = Path(target)

    search_candidates: list[Path] = []
    if raw_candidate.is_absolute():
        search_candidates.append(raw_candidate)
    else:
        search_candidates.extend(
            [
                md_dir / raw_candidate,
                md_dir / raw_candidate.name,
                md_dir / "charts" / raw_candidate.name,
                md_dir / "images" / raw_candidate.name,
                md_dir / "figures" / raw_candidate.name,
            ]
        )

    for candidate in search_candidates:
        try:
            resolved = candidate.resolve()
        except Exception:
            resolved = candidate.absolute()
        if resolved.exists() and resolved.is_file():
            return resolved

    if not raw_candidate.is_absolute() and raw_candidate.name:
        for folder_name in ["charts", "images", "figures"]:
            folder = md_dir / folder_name
            if not folder.exists() or not folder.is_dir():
                continue
            matches = list(folder.rglob(raw_candidate.name))
            if matches:
                return matches[0]

    return None


def get_current_files(work_dir: str, sub_dir: str = "") -> str:
    """获取当前工作目录下的文件列表"""
    path = Path(work_dir)
    if sub_dir:
        path = path / sub_dir

    if not path.exists():
        return "目录不存在"

    files = []
    for f in sorted(path.iterdir()):
        if f.is_file():
            size_kb = f.stat().st_size / 1024
            display_name = f"{sub_dir}/{f.name}" if sub_dir else f.name
            files.append(f"- {display_name} ({size_kb:.1f} KB)")

    if not files:
        return "目录为空"
    return "\n".join(files)


def read_file_content(filepath: str) -> Optional[str]:
    """读取文件内容"""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        return None


def save_json(data: dict, filepath: str) -> None:
    """保存 JSON 文件"""
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_json(filepath: str) -> Optional[dict]:
    """加载 JSON 文件"""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def normalize_math_markdown(markdown_text: str) -> str:
    """规范化 Markdown 中的 LaTeX 公式分隔符，便于前端渲染与 DOCX 导出。"""
    if not markdown_text:
        return markdown_text

    def normalize_math_body(math_body: str) -> str:
        body = math_body.strip()
        body = re.sub(r"\\\\(?=[A-Za-z])", r"\\", body)
        return body

    code_fence_pattern = re.compile(r"(```[\s\S]*?```)")
    segments = code_fence_pattern.split(markdown_text.replace("\r\n", "\n"))
    normalized_segments: list[str] = []

    for segment in segments:
        if segment.startswith("```"):
            normalized_segments.append(segment)
            continue

        segment = (
            segment
            .replace("\\\\[", "\\[")
            .replace("\\\\]", "\\]")
            .replace("\\\\(", "\\(")
            .replace("\\\\)", "\\)")
        )

        segment = re.sub(
            r"\\\[\s*([\s\S]*?)\s*\\\]",
            lambda m: f"\n$$\n{normalize_math_body(m.group(1))}\n$$\n",
            segment,
        )
        segment = re.sub(
            r"\\\(\s*([\s\S]*?)\s*\\\)",
            lambda m: f"${normalize_math_body(m.group(1))}$",
            segment,
        )
        segment = re.sub(
            r"\$\$\s*([\s\S]*?)\s*\$\$",
            lambda m: f"$$\n{normalize_math_body(m.group(1))}\n$$",
            segment,
        )
        segment = re.sub(
            r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)",
            lambda m: f"${normalize_math_body(m.group(1))}$",
            segment,
            flags=re.S,
        )
        segment = re.sub(r"\n{3,}", "\n\n", segment)
        normalized_segments.append(segment)

    return "".join(normalized_segments).strip() + "\n"


def extract_markdown_formulas(markdown_text: str) -> list[dict[str, Any]]:
    """抽取 Markdown 中的内联与块级公式，供终审复核使用。"""
    text = markdown_text or ""
    formulas: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str]] = set()

    patterns = [
        ("block", re.compile(r"\$\$\s*([\s\S]*?)\s*\$\$")),
        ("block", re.compile(r"\\\[\s*([\s\S]*?)\s*\\\]")),
        ("inline", re.compile(r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)", re.S)),
        ("inline", re.compile(r"\\\((.+?)\\\)", re.S)),
    ]

    for formula_type, pattern in patterns:
        for match in pattern.finditer(text):
            formula = " ".join((match.group(1) or "").split())
            if not formula:
                continue
            start = match.start()
            line_number = text.count("\n", 0, start) + 1
            line_start = text.rfind("\n", 0, start)
            line_end = text.find("\n", match.end())
            snippet_start = 0 if line_start < 0 else line_start + 1
            snippet_end = len(text) if line_end < 0 else line_end
            raw_line = text[snippet_start:snippet_end].strip()
            number_match = re.search(r"[（(]\s*(\d+)\s*[)）]", raw_line)
            equation_no = number_match.group(1) if number_match else ""
            key = (formula_type, formula, equation_no)
            if key in seen:
                continue
            seen.add(key)
            formulas.append(
                {
                    "type": formula_type,
                    "line": line_number,
                    "equation_no": equation_no,
                    "expression": formula,
                    "raw_context": raw_line,
                }
            )

    return formulas


def extract_markdown_tables(markdown_text: str) -> list[dict[str, Any]]:
    """抽取 Markdown 表格及其中的数值，供终审复核使用。"""
    lines = (markdown_text or "").splitlines()
    tables: list[dict[str, Any]] = []
    numeric_pattern = re.compile(r"[-+]?\d+(?:\.\d+)?(?:[eE][-+]?\d+)?%?")
    index = 1
    current_block: list[tuple[int, str]] = []

    def flush_table() -> None:
        nonlocal current_block, index
        if len(current_block) < 2:
            current_block = []
            return

        stripped_lines = [line.strip() for _, line in current_block if line.strip()]
        if len(stripped_lines) < 2:
            current_block = []
            return

        separator = stripped_lines[1]
        if not re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", separator):
            current_block = []
            return

        header = [cell.strip() for cell in stripped_lines[0].strip("|").split("|")]
        rows: list[list[str]] = []
        numeric_values: list[dict[str, Any]] = []

        for row_offset, raw in enumerate(stripped_lines[2:], start=1):
            row = [cell.strip() for cell in raw.strip("|").split("|")]
            rows.append(row)
            for col_index, cell in enumerate(row):
                values = numeric_pattern.findall(cell)
                if not values:
                    continue
                numeric_values.append(
                    {
                        "row": row_offset,
                        "column": header[col_index] if col_index < len(header) else f"col_{col_index + 1}",
                        "cell": cell,
                        "values": values,
                    }
                )

        tables.append(
            {
                "table_index": index,
                "start_line": current_block[0][0],
                "headers": header,
                "rows": rows,
                "numeric_values": numeric_values,
            }
        )
        index += 1
        current_block = []

    for line_number, line in enumerate(lines, start=1):
        if "|" in line and line.strip().startswith("|"):
            current_block.append((line_number, line))
        else:
            flush_table()
    flush_table()
    return tables


def extract_reference_entries(markdown_text: str) -> list[str]:
    """抽取参考文献段落或引用条目。"""
    lines = (markdown_text or "").splitlines()
    entries: list[str] = []
    in_reference_section = False

    for raw_line in lines:
        line = raw_line.strip()
        if re.match(r"^#{1,6}\s*(参考文献|References)\s*$", line, re.I):
            in_reference_section = True
            continue

        if in_reference_section:
            if re.match(r"^#{1,6}\s+", line):
                break
            if line:
                entries.append(line)
            continue

        if re.match(r"^\[(\d+|[A-Za-z]+)\]\s+", line):
            entries.append(line)

    return entries


def extract_problem_parameters(question_text: str) -> list[dict[str, Any]]:
    """从原题中抽取带数值的参数语句，供终审检查“题设是否被擅自修改”。"""
    numeric_pattern = re.compile(r"[-+]?\d+(?:\.\d+)?(?:[eE][-+]?\d+)?%?")
    parameters: list[dict[str, Any]] = []

    for line_number, raw_line in enumerate((question_text or "").splitlines(), start=1):
        line = raw_line.strip()
        if not line:
            continue
        numbers = numeric_pattern.findall(line)
        if not numbers:
            continue
        parameters.append(
            {
                "line": line_number,
                "text": line,
                "numbers": numbers,
            }
        )

    return parameters


def build_paper_audit_manifest(question_text: str, markdown_text: str) -> dict[str, Any]:
    """构建终审所需的结构化抽取清单。"""
    formulas = extract_markdown_formulas(markdown_text)
    tables = extract_markdown_tables(markdown_text)
    references = extract_reference_entries(markdown_text)
    problem_parameters = extract_problem_parameters(question_text)

    placeholder_markers = []
    for marker in ["TODO", "TBD", "待补", "占位", "xx", "xxx", "et al.", "作者待补", "年份待补"]:
        if marker.lower() in (markdown_text or "").lower():
            placeholder_markers.append(marker)

    return {
        "formula_count": len(formulas),
        "table_count": len(tables),
        "reference_count": len(references),
        "problem_parameter_count": len(problem_parameters),
        "formulas": formulas,
        "tables": tables,
        "references": references,
        "problem_parameters": problem_parameters,
        "reference_placeholder_markers": placeholder_markers,
    }


def md_to_docx(md_path: str, docx_path: str) -> bool:
    """尝试将 markdown 转为 docx，优先保留 LaTeX 公式为 Word 数学对象。"""
    try:
        import pypandoc

        pypandoc.convert_file(
            md_path,
            "docx",
            outputfile=docx_path,
            extra_args=[
                f"--resource-path={Path(md_path).parent}{os.pathsep}{Path(md_path).parent / 'charts'}"
            ],
        )
        # pypandoc 在 outputfile 模式下通常返回空字符串
        return os.path.exists(docx_path)
    except Exception:
        pass

    try:
        from docx import Document
        from docx.shared import Inches
        from docx.shared import Pt
    except Exception:
        return False

    try:
        markdown = normalize_math_markdown(Path(md_path).read_text(encoding="utf-8"))
    except Exception:
        return False

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    lines = markdown.splitlines()
    in_code_block = False
    code_buffer: list[str] = []
    table_buffer: list[str] = []

    def flush_code() -> None:
        nonlocal code_buffer
        if not code_buffer:
            return
        paragraph = document.add_paragraph()
        run = paragraph.add_run("\n".join(code_buffer))
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        code_buffer = []

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = []
        for raw in table_buffer:
            stripped = raw.strip()
            if not stripped:
                continue
            if re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", stripped):
                continue
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            rows.append(cells)
        if rows:
            column_count = max(len(row) for row in rows)
            table = document.add_table(rows=len(rows), cols=column_count)
            table.style = "Table Grid"
            for row_index, row in enumerate(rows):
                for col_index, cell in enumerate(row):
                    table.cell(row_index, col_index).text = cell
        table_buffer = []

    for line in lines:
        stripped = line.rstrip()
        if stripped.startswith("```"):
            flush_table()
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(stripped)
            continue

        if "|" in stripped and stripped.count("|") >= 2:
            table_buffer.append(stripped)
            continue

        flush_table()

        if not stripped.strip():
            document.add_paragraph("")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 6)
            document.add_heading(heading_match.group(2).strip(), level=level)
            continue

        bullet_match = re.match(r"^\s*[-*+]\s+(.*)$", stripped)
        if bullet_match:
            document.add_paragraph(bullet_match.group(1).strip(), style="List Bullet")
            continue

        ordered_match = re.match(r"^\s*\d+[.)]\s+(.*)$", stripped)
        if ordered_match:
            document.add_paragraph(ordered_match.group(1).strip(), style="List Number")
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped.strip())
        if image_match:
            alt_text = image_match.group(1).strip() or "图片"
            path_text = image_match.group(2).strip()
            image_path = _resolve_markdown_image_path(md_path, path_text)
            if image_path:
                picture_paragraph = document.add_paragraph()
                picture_paragraph.alignment = 1
                try:
                    picture_paragraph.add_run().add_picture(str(image_path), width=Inches(6))
                except Exception:
                    picture_paragraph.add_run().add_picture(str(image_path))

                caption = document.add_paragraph(alt_text)
                caption.alignment = 1
            else:
                document.add_paragraph(f"[图片缺失] {alt_text}: {path_text}")
            continue

        document.add_paragraph(stripped)

    flush_table()
    flush_code()

    try:
        document.save(docx_path)
    except Exception:
        return False
    return os.path.exists(docx_path)
